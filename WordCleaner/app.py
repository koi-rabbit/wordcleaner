# app.py
import streamlit as st
from docx import Document
import re, os
from io import BytesIO
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.shared import Cm

# 页面配置
st.set_page_config(
    page_title="Word自动排版工具",
    page_icon="📝",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ========== 初始化 session_state ==========
if 'params_initialized' not in st.session_state:
    # 默认参数配置
    defaults = {
        # 正文
        "bdy_cz_font_name": "宋体",
        "bdy_font_name": "Times New Roman",
        "bdy_font_size": 10.5,
        "bdy_space_before": 6.0,
        "bdy_space_after": 6.0,
        "bdy_line_spacing": 1.0,
        "bdy_first_line_indent": 0.75,
        
        # 表格
        "tbl_cz_font_name": "宋体",
        "tbl_font_name": "Times New Roman",
        "tbl_font_size": 10.5,
        "tbl_space_before": 4.0,
        "tbl_space_after": 4.0,
        "tbl_line_spacing": 1.0,
        "tbl_width": 6.0,
        
        # 标题样式 (1-9级)
        "h1_cz_font": "黑体",
        "h1_font": "Arial",
        "h1_size": 14,
        "h1_bold": False,
        "h1_before": 12,
        "h1_after": 12,
        "h1_line": 1.5,
        "h1_indent": 0,
        
        "h2_cz_font": "黑体",
        "h2_font": "Arial",
        "h2_size": 12,
        "h2_bold": False,
        "h2_before": 12,
        "h2_after": 12,
        "h2_line": 1.5,
        "h2_indent": 0.75,
        
        "h3_cz_font": "宋体",
        "h3_font": "Times New Roman",
        "h3_size": 10.5,
        "h3_bold": False,
        "h3_before": 8,
        "h3_after": 8,
        "h3_line": 1.0,
        "h3_indent": 1.5,
        
        "h4_cz_font": "宋体",
        "h4_font": "Times New Roman",
        "h4_size": 10.5,
        "h4_bold": False,
        "h4_before": 8,
        "h4_after": 8,
        "h4_line": 1.0,
        "h4_indent": 2.25,
        
        "h5_cz_font": "宋体",
        "h5_font": "Times New Roman",
        "h5_size": 10.5,
        "h5_bold": False,
        "h5_before": 6,
        "h5_after": 6,
        "h5_line": 1.0,
        "h5_indent": 3.0,
        
        "h6_cz_font": "宋体",
        "h6_font": "Arial",
        "h6_size": 9,
        "h6_bold": False,
        "h6_before": 2,
        "h6_after": 2,
        "h6_line": 1.0,
        "h6_indent": 0,
        
        "h7_cz_font": "宋体",
        "h7_font": "Arial",
        "h7_size": 8,
        "h7_bold": False,
        "h7_before": 0,
        "h7_after": 0,
        "h7_line": 1.0,
        "h7_indent": 0,
        
        "h8_cz_font": "宋体",
        "h8_font": "Arial",
        "h8_size": 7,
        "h8_bold": False,
        "h8_before": 0,
        "h8_after": 0,
        "h8_line": 1.0,
        "h8_indent": 0,
        
        "h9_cz_font": "宋体",
        "h9_font": "Arial",
        "h9_size": 6,
        "h9_bold": False,
        "h9_before": 0,
        "h9_after": 0,
        "h9_line": 1.0,
        "h9_indent": 0,
        
        # 编号方案选择
        "numbering_scheme": "方案一：中文数字",
    }
    
    st.session_state.update(defaults)
    st.session_state['params_initialized'] = True

# ========== 侧边栏：简洁的参数设置 ==========
with st.sidebar:
    st.title("⚙️ 排版设置")
    
    # 使用选项卡组织设置
    tab1, tab2, tab3 = st.tabs(["标题", "正文", "表格"])
    
    with tab1:
        # 标题级别选择器
        heading_level = st.selectbox(
            "选择标题级别",
            options=["1级", "2级", "3级", "4级", "5级", "6级", "7级", "8级", "9级"],
            index=0
        )
        
        # 获取当前级别对应的键前缀
        level_num = int(heading_level[0])
        prefix = f"h{level_num}_"
        
        # 编号方案和样式显示 - 放到一行
        col_scheme, col_style = st.columns([2, 3])
        with col_scheme:
            st.session_state["numbering_scheme"] = st.selectbox(
                "编号方案",
                options=["方案一：中文数字", "方案二：数字点号"],
                index=0 if st.session_state["numbering_scheme"] == "方案一：中文数字" else 1,
                key="numbering_scheme_select",
                label_visibility="collapsed"
            )
        
        with col_style:
            # 根据选择的方案显示对应级别的编号样式
            if st.session_state["numbering_scheme"] == "方案一：中文数字":
                scheme_styles = {
                    1: "一、", 2: "（一）", 3: "1.", 4: "（1）", 
                    5: "（i）", 6: "（a）", 7: "（一）", 8: "（1）", 9: "（i）"
                }
            else:
                # 方案二：数字点号方案
                scheme_styles = {
                    1: "1.", 2: "1.1", 3: "1.1.1", 4: "1.1.1.1", 
                    5: "1.1.1.1.1", 6: "1.1.1.1.1.1", 
                    7: "1.1.1.1.1.1.1", 8: "1.1.1.1.1.1.1.1", 
                    9: "1.1.1.1.1.1.1.1.1"
                }
            
            current_style = scheme_styles[level_num]
            st.text_input(
                "编号样式",
                value=current_style,
                disabled=True,
                key=f"number_style_display_{level_num}",
                label_visibility="collapsed"
            )
        
        st.markdown("---")
        
        # 字体设置
        col1, col2 = st.columns(2)
        with col1:
            st.session_state[f"{prefix}cz_font"] = st.selectbox(
                "中文字体",
                ["黑体", "宋体", "楷体", "仿宋", "微软雅黑"],
                index=["黑体", "宋体", "楷体", "仿宋", "微软雅黑"].index(st.session_state[f"{prefix}cz_font"]),
                key=f"{prefix}cz_font_select"
            )
        with col2:
            st.session_state[f"{prefix}font"] = st.selectbox(
                "英文字体",
                ["Arial", "Times New Roman", "Calibri", "Verdana"],
                index=["Arial", "Times New Roman", "Calibri", "Verdana"].index(st.session_state[f"{prefix}font"]),
                key=f"{prefix}font_select"
            )
        
        # 字体大小和粗体设置
        col_size, col_bold = st.columns(2)
        with col_size:
            st.session_state[f"{prefix}size"] = st.number_input(
                "字体大小(pt)",
                min_value=6.0,
                max_value=20.0,
                value=float(st.session_state[f"{prefix}size"]),
                step=0.5,
                key=f"{prefix}size_input"
            )
        with col_bold:
            # 粗体改为下拉选择
            bold_options = ["否", "是"]
            current_bold = "是" if st.session_state[f"{prefix}bold"] else "否"
            selected_bold = st.selectbox(
                "粗体",
                options=bold_options,
                index=bold_options.index(current_bold),
                key=f"{prefix}bold_select"
            )
            st.session_state[f"{prefix}bold"] = (selected_bold == "是")
        
        # 间距设置
        col_before, col_after = st.columns(2)
        with col_before:
            st.session_state[f"{prefix}before"] = st.number_input(
                "段前间距(pt)",
                min_value=0.0,
                max_value=20.0,
                value=float(st.session_state[f"{prefix}before"]),
                step=0.5,
                key=f"{prefix}before_input"
            )
        with col_after:
            st.session_state[f"{prefix}after"] = st.number_input(
                "段后间距(pt)",
                min_value=0.0,
                max_value=20.0,
                value=float(st.session_state[f"{prefix}after"]),
                step=0.5,
                key=f"{prefix}after_input"
            )
        
        # 行距和缩进
        col_line, col_indent = st.columns(2)
        with col_line:
            st.session_state[f"{prefix}line"] = st.number_input(
                "行间距(倍)",
                min_value=1.0,
                max_value=3.0,
                value=float(st.session_state[f"{prefix}line"]),
                step=0.1,
                key=f"{prefix}line_input"
            )
        with col_indent:
            st.session_state[f"{prefix}indent"] = st.number_input(
                "首行缩进(cm)",
                min_value=0.0,
                max_value=5.0,
                value=float(st.session_state[f"{prefix}indent"]),
                step=0.1,
                key=f"{prefix}indent_input"
            )
        
        # 重置当前标题级别的按钮
        st.markdown("---")
        if st.button(f"🔄 重置{heading_level}设置", use_container_width=True):
            # 重置当前级别的设置
            reset_keys = {
                f"{prefix}cz_font": defaults[f"{prefix}cz_font"],
                f"{prefix}font": defaults[f"{prefix}font"],
                f"{prefix}size": defaults[f"{prefix}size"],
                f"{prefix}bold": defaults[f"{prefix}bold"],
                f"{prefix}before": defaults[f"{prefix}before"],
                f"{prefix}after": defaults[f"{prefix}after"],
                f"{prefix}line": defaults[f"{prefix}line"],
                f"{prefix}indent": defaults[f"{prefix}indent"],
            }
            for key, value in reset_keys.items():
                st.session_state[key] = value
            st.success(f"{heading_level}设置已重置！")
            st.rerun()
    
    with tab2:
        # 正文格式设置
        st.markdown("### 正文格式")
        
        # 字体设置
        col_bdy_font1, col_bdy_font2 = st.columns(2)
        with col_bdy_font1:
            st.session_state["bdy_cz_font_name"] = st.selectbox(
                "中文字体",
                ["宋体", "黑体", "楷体", "仿宋", "微软雅黑"],
                index=["宋体", "黑体", "楷体", "仿宋", "微软雅黑"].index(st.session_state["bdy_cz_font_name"]),
                key="bdy_cz_font_select"
            )
        with col_bdy_font2:
            st.session_state["bdy_font_name"] = st.selectbox(
                "英文字体",
                ["Times New Roman", "Arial", "Calibri", "Verdana"],
                index=["Times New Roman", "Arial", "Calibri", "Verdana"].index(st.session_state["bdy_font_name"]),
                key="bdy_font_select"
            )
        
        # 字体大小
        st.session_state["bdy_font_size"] = st.number_input(
            "字体大小(pt)",
            min_value=8.0,
            max_value=16.0,
            value=float(st.session_state["bdy_font_size"]),
            step=0.5,
            key="bdy_size_input"
        )
        
        # 间距设置
        col_bdy_before, col_bdy_after = st.columns(2)
        with col_bdy_before:
            st.session_state["bdy_space_before"] = st.number_input(
                "段前间距(pt)",
                min_value=0.0,
                max_value=20.0,
                value=float(st.session_state["bdy_space_before"]),
                step=0.5,
                key="bdy_before_input"
            )
        with col_bdy_after:
            st.session_state["bdy_space_after"] = st.number_input(
                "段后间距(pt)",
                min_value=0.0,
                max_value=20.0,
                value=float(st.session_state["bdy_space_after"]),
                step=0.5,
                key="bdy_after_input"
            )
        
        # 行距和缩进
        col_bdy_line, col_bdy_indent = st.columns(2)
        with col_bdy_line:
            st.session_state["bdy_line_spacing"] = st.number_input(
                "行间距(倍)",
                min_value=0.5,
                max_value=3.0,
                value=float(st.session_state["bdy_line_spacing"]),
                step=0.1,
                key="bdy_line_input"
            )
        with col_bdy_indent:
            st.session_state["bdy_first_line_indent"] = st.number_input(
                "首行缩进(cm)",
                min_value=0.0,
                max_value=2.0,
                value=float(st.session_state["bdy_first_line_indent"]),
                step=0.1,
                key="bdy_indent_input"
            )
        
        # 重置正文设置按钮
        st.markdown("---")
        if st.button("🔄 重置正文设置", use_container_width=True):
            reset_keys = {
                "bdy_cz_font_name": defaults["bdy_cz_font_name"],
                "bdy_font_name": defaults["bdy_font_name"],
                "bdy_font_size": defaults["bdy_font_size"],
                "bdy_space_before": defaults["bdy_space_before"],
                "bdy_space_after": defaults["bdy_space_after"],
                "bdy_line_spacing": defaults["bdy_line_spacing"],
                "bdy_first_line_indent": defaults["bdy_first_line_indent"],
            }
            for key, value in reset_keys.items():
                st.session_state[key] = value
            st.success("正文设置已重置！")
            st.rerun()
    
    with tab3:
        # 表格格式设置
        st.markdown("### 表格格式")
        
        # 字体设置
        col_tbl_font1, col_tbl_font2 = st.columns(2)
        with col_tbl_font1:
            st.session_state["tbl_cz_font_name"] = st.selectbox(
                "中文字体",
                ["宋体", "黑体", "楷体", "仿宋", "微软雅黑"],
                index=["宋体", "黑体", "楷体", "仿宋", "微软雅黑"].index(st.session_state["tbl_cz_font_name"]),
                key="tbl_cz_font_select"
            )
        with col_tbl_font2:
            st.session_state["tbl_font_name"] = st.selectbox(
                "英文字体",
                ["Times New Roman", "Arial", "Calibri", "Verdana"],
                index=["Times New Roman", "Arial", "Calibri", "Verdana"].index(st.session_state["tbl_font_name"]),
                key="tbl_font_select"
            )
        
        # 字体大小
        st.session_state["tbl_font_size"] = st.number_input(
            "字体大小(pt)",
            min_value=8.0,
            max_value=14.0,
            value=float(st.session_state["tbl_font_size"]),
            step=0.5,
            key="tbl_size_input"
        )
        
        # 间距设置
        col_tbl_before, col_tbl_after = st.columns(2)
        with col_tbl_before:
            st.session_state["tbl_space_before"] = st.number_input(
                "段前间距(pt)",
                min_value=0.0,
                max_value=10.0,
                value=float(st.session_state["tbl_space_before"]),
                step=0.5,
                key="tbl_before_input"
            )
        with col_tbl_after:
            st.session_state["tbl_space_after"] = st.number_input(
                "段后间距(pt)",
                min_value=0.0,
                max_value=10.0,
                value=float(st.session_state["tbl_space_after"]),
                step=0.5,
                key="tbl_after_input"
            )
        
        # 行距和宽度
        col_tbl_line, col_tbl_width = st.columns(2)
        with col_tbl_line:
            st.session_state["tbl_line_spacing"] = st.number_input(
                "行间距(倍)",
                min_value=0.5,
                max_value=3.0,
                value=float(st.session_state["tbl_line_spacing"]),
                step=0.1,
                key="tbl_line_input"
            )
        with col_tbl_width:
            st.session_state["tbl_width"] = st.number_input(
                "表格宽度(英寸)",
                min_value=3.0,
                max_value=10.0,
                value=float(st.session_state["tbl_width"]),
                step=0.1,
                key="tbl_width_input"
            )
        
        # 重置表格设置按钮
        st.markdown("---")
        if st.button("🔄 重置表格设置", use_container_width=True):
            reset_keys = {
                "tbl_cz_font_name": defaults["tbl_cz_font_name"],
                "tbl_font_name": defaults["tbl_font_name"],
                "tbl_font_size": defaults["tbl_font_size"],
                "tbl_space_before": defaults["tbl_space_before"],
                "tbl_space_after": defaults["tbl_space_after"],
                "tbl_line_spacing": defaults["tbl_line_spacing"],
                "tbl_width": defaults["tbl_width"],
            }
            for key, value in reset_keys.items():
                st.session_state[key] = value
            st.success("表格设置已重置！")
            st.rerun()

# ========== 主页面：简洁的文件处理界面 ==========
st.title("📝 Word自动排版工具")
st.markdown("---")

# 简介
st.markdown("""
**快速开始：**
1. 📌 **设置格式** - 在左侧选择"标题"、"正文"或"表格"选项卡，调整对应格式参数
2. 📤 **上传文档** - 支持批量上传多个Word文档
3. 🚀 **开始处理** - 点击下方"开始处理文档"按钮
4. 📥 **下载结果** - 处理完成后下载排版后的文档

**功能特点：**
- 🎯 **智能排版**：自动识别文档大纲结构
- 🔢 **自动编号**：智能添加多级标题序号
- 🎨 **格式统一**：批量设置文档格式
- ⚡ **高效处理**：支持多文件同时处理
""")

# 文件上传区域
st.markdown("### 📤 文档上传")
uploaded_files = st.file_uploader(
    "选择Word文档 (.docx)",
    type=["docx"],
    accept_multiple_files=True,
    help="支持批量上传多个文档",
    label_visibility="collapsed"
)

# 显示已上传文件
if uploaded_files:
    st.success(f"✅ 已选择 {len(uploaded_files)} 个文档")
    
    # 文件列表
    with st.expander("📋 文件列表", expanded=True):
        for i, file in enumerate(uploaded_files, 1):
            col1, col2, col3 = st.columns([6, 2, 2])
            with col1:
                st.write(f"**{file.name}**")
            with col2:
                st.write(f"`{file.size / 1024:.1f} KB`")
            with col3:
                st.write("📄")
    
    # 处理按钮
    st.markdown("---")
    
    # 处理选项
    col1, col2 = st.columns(2)
    with col1:
        add_numbers = st.checkbox("添加标题序号", value=True, help="自动为标题添加层级序号")
    with col2:
        keep_format = st.checkbox("保留原有格式", value=False, help="尽量保留文档原有格式")
    
    # 处理按钮
    if st.button("🚀 开始处理文档", type="primary", use_container_width=True):
        # 创建进度条
        progress_bar = st.progress(0)
        status_text = st.empty()
        
        # 处理结果区域
        results_container = st.container()
        
        # 从session_state获取当前样式配置
        style_rules = {}
        for level in range(1, 10):
            prefix = f"h{level}_"
            style_rules[level] = {
                'style_name': f'Heading {level}',
                'cz_font_name': st.session_state[f"{prefix}cz_font"],
                'font_name': st.session_state[f"{prefix}font"],
                'font_size': st.session_state[f"{prefix}size"],
                'bold': st.session_state[f"{prefix}bold"],
                'space_before': st.session_state[f"{prefix}before"],
                'space_after': st.session_state[f"{prefix}after"],
                'line_spacing': st.session_state[f"{prefix}line"],
                'first_line_indent': st.session_state[f"{prefix}indent"],
            }
        
        # 获取编号方案
        numbering_scheme = st.session_state["numbering_scheme"]
        
        # 处理每个文件
        with results_container:
            for idx, uploaded_file in enumerate(uploaded_files):
                # 更新进度
                progress = (idx + 1) / len(uploaded_files)
                progress_bar.progress(progress)
                status_text.text(f"正在处理: **{uploaded_file.name}** ({idx + 1}/{len(uploaded_files)})")
                
                try:
                    # 处理文档
                    processed_buffer = process_single_document(
                        uploaded_file.read(),
                        style_rules,
                        numbering_scheme,
                        st.session_state,
                        add_numbers
                    )
                    
                    # 显示处理结果
                    col_result1, col_result2, col_result3 = st.columns([6, 3, 1])
                    with col_result1:
                        st.write(f"✅ **{uploaded_file.name}**")
                    with col_result2:
                        st.download_button(
                            label="📥 下载文件",
                            data=processed_buffer,
                            file_name=f"排版_{uploaded_file.name}",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key=f"download_{idx}",
                            use_container_width=True
                        )
                    
                except Exception as e:
                    st.error(f"❌ 处理 {uploaded_file.name} 时出错: `{str(e)}`")
            
            # 完成提示
            progress_bar.empty()
            status_text.success("✅ 所有文档处理完成！")
            st.balloons()

else:
    st.info("👈 请先在左侧设置参数，然后上传需要排版的Word文档")

# ========== 工具函数定义 ==========
KNOWN_STYLES = {
    "Normal",
    "List Paragraph",
    "Heading 1", "Heading 2", "Heading 3", "Heading 4",
    "Heading 5", "Heading 6", "Heading 7", "Heading 8", "Heading 9"
}

def get_outline_level_from_xml(p):
    """从段落的XML中提取大纲级别，并加1"""
    xml = p._p.xml
    m = re.search(r'<w:outlineLvl w:val="(\d)"/>', xml)
    level = int(m.group(1)) if m else None
    if level is not None:
        level += 1
    return level

def restructure_outline(doc):
    """重构文档大纲"""
    for p in doc.paragraphs:
        zero_indent(p)
        lvl = get_outline_level_from_xml(p)
        if lvl and p.style.name == "Normal":
            heading_style = f"Heading {lvl}"
            if heading_style in doc.styles:
                p.style = doc.styles[heading_style]
    
    # 降级空标题
    for p in doc.paragraphs:
        if p.style.name.startswith("Heading") and not p.text.strip():
            p.style = doc.styles["Normal"]

def zero_indent(p):
    """清除段落缩进"""
    pf = p.paragraph_format
    pf.left_indent = Cm(0)
    pf.first_line_indent = Cm(0)
    pf.right_indent = Cm(0)
    pf.tab_stops.clear_all()
    if p.text:
        p.text = p.text.lstrip()

def kill_all_numbering(doc):
    """清除所有编号"""
    for st_name in ['List Paragraph', 'Heading 1', 'Heading 2', 'Heading 3',
                    'Heading 4', 'Heading 5', 'Heading 6', 'Heading 7',
                    'Heading 8', 'Heading 9']:
        try:
            style = doc.styles[st_name]
        except KeyError:
            continue
        style_el = style._element
        for num_id in style_el.xpath('.//w:numId'):
            num_id.getparent().remove(num_id)

def set_font(run, cz_font_name, font_name):
    """设置字体"""
    rPr = run.element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:eastAsia'), cz_font_name)
    rFonts.set(qn('w:ascii'), font_name)

def number_to_chinese(number):
    """数字转中文大写数字"""
    chinese_numbers = ["零", "一", "二", "三", "四", "五", "六", "七", "八", "九", "十"]
    
    if number <= 10:
        return chinese_numbers[number]
    elif number < 20:
        return "十" + (chinese_numbers[number - 10] if number != 10 else "")
    elif number < 100:
        tens = number // 10
        ones = number % 10
        if ones == 0:
            return chinese_numbers[tens] + "十"
        else:
            return chinese_numbers[tens] + "十" + chinese_numbers[ones]
    else:
        return str(number)

def number_to_roman(num):
    """将数字转换为罗马数字（1-20）"""
    roman_numerals = {
        1: 'i', 2: 'ii', 3: 'iii', 4: 'iv', 5: 'v',
        6: 'vi', 7: 'vii', 8: 'viii', 9: 'ix', 10: 'x',
        11: 'xi', 12: 'xii', 13: 'xiii', 14: 'xiv', 15: 'xv',
        16: 'xvi', 17: 'xvii', 18: 'xviii', 19: 'xix', 20: 'xx'
    }
    return roman_numerals.get(num, str(num))

def number_to_letter(num):
    """将数字转换为字母（1-26）"""
    if 1 <= num <= 26:
        return chr(96 + num)  # a-z
    return str(num)

def add_heading_numbers_custom(doc, numbering_scheme, add_numbers=True):
    """添加自定义标题序号"""
    if not add_numbers:
        return
    
    number_pattern = re.compile(
        r'^\s*'
        r'[（(]?'
        r'[\d一二三四五六七八九十零①②③④⑤⑥⑦⑧⑨⑩]{1,4}'
        r'[\.、）)\s]'
        r'(?:[\d一二三四五六七八九十零①②③④⑤⑥⑦⑧⑨⑩]{1,4}'
        r'[\.、）)\s]'
        r')*',
        re.UNICODE
    )
    
    heading_numbers = [0] * 9
    
    for paragraph in doc.paragraphs:
        if paragraph.style.name.startswith('Heading'):
            if paragraph.text == "Ellipsis" or not paragraph.text.strip():
                continue
            
            # 清除原有编号
            for p in doc.paragraphs:
                p_pr = p._p.get_or_add_pPr()
                num_pr = p_pr.find(qn('w:numPr'))
                if num_pr is not None:
                    p_pr.remove(num_pr)
            
            paragraph.text = number_pattern.sub('', paragraph.text).strip()
            level = int(paragraph.style.name.split(' ')[1]) - 1
            
            # 更新序号
            heading_numbers[level] += 1
            for i in range(level + 1, len(heading_numbers)):
                heading_numbers[i] = 0
            
            # 添加序号
            if heading_numbers[level] > 0:
                if numbering_scheme == "方案一：中文数字":
                    # 方案一：中文数字方案
                    number_str = get_scheme1_number(level + 1, heading_numbers)
                else:
                    # 方案二：数字点号方案
                    number_str = get_scheme2_number(level + 1, heading_numbers)
                
                paragraph.text = number_str + paragraph.text

def get_scheme1_number(level, heading_numbers):
    """获取方案一的编号"""
    if level == 1:
        # 一、
        return number_to_chinese(heading_numbers[0]) + "、"
    elif level == 2:
        # （一）
        return "（" + number_to_chinese(heading_numbers[1]) + "）"
    elif level == 3:
        # 1.
        return str(heading_numbers[2]) + "."
    elif level == 4:
        # （1）
        return "（" + str(heading_numbers[3]) + "）"
    elif level == 5:
        # （i）
        return "（" + number_to_roman(heading_numbers[4]) + "）"
    elif level == 6:
        # （a）
        return "（" + number_to_letter(heading_numbers[5]) + "）"
    elif level == 7:
        # （一）
        return "（" + number_to_chinese(heading_numbers[6]) + "）"
    elif level == 8:
        # （1）
        return "（" + str(heading_numbers[7]) + "）"
    elif level == 9:
        # （i）
        return "（" + number_to_roman(heading_numbers[8]) + "）"
    return str(heading_numbers[level-1]) + "."

def get_scheme2_number(level, heading_numbers):
    """获取方案二的编号（1.1.1格式）"""
    # 只取当前级别及以上的数字
    numbers = heading_numbers[:level]
    # 连接成 1.1.1 格式
    return '.'.join(str(num) for num in numbers if num > 0) + "."

def process_single_document(file_bytes, style_rules, numbering_scheme, params, add_numbers=True):
    """处理单个文档"""
    doc = Document(BytesIO(file_bytes))
    
    # 重构大纲
    restructure_outline(doc)
    
    # 清除编号
    kill_all_numbering(doc)
    
    # 添加标题序号
    add_heading_numbers_custom(doc, numbering_scheme, add_numbers)
    
    # 应用格式
    skipped = set()
    
    for p in doc.paragraphs:
        style_name = p.style.name
        
        if p.text == "Ellipsis" or not p.text.strip():
            continue
        
        if style_name not in KNOWN_STYLES:
            skipped.add(style_name)
            continue
        
        if style_name.startswith("Heading"):
            level = int(style_name.split(' ')[1])
            rule = style_rules[level]
            p.style.paragraph_format.space_before = Pt(rule['space_before'])
            p.style.paragraph_format.space_after = Pt(rule['space_after'])
            p.style.paragraph_format.line_spacing = rule['line_spacing']
            p.style.paragraph_format.first_line_indent = Cm(rule['first_line_indent'])
            for run in p.runs:
                set_font(run, rule['cz_font_name'], rule['font_name'])
                run.font.size = Pt(rule['font_size'])
                run.font.bold = rule['bold']
        else:
            # 正文格式
            p.paragraph_format.space_before = Pt(params['bdy_space_before'])
            p.paragraph_format.space_after = Pt(params['bdy_space_after'])
            p.paragraph_format.line_spacing = params['bdy_line_spacing']
            p.paragraph_format.first_line_indent = Cm(params['bdy_first_line_indent'])
            for run in p.runs:
                set_font(run, params['bdy_cz_font_name'], params['bdy_font_name'])
                run.font.size = Pt(params['bdy_font_size'])
    
    # 表格格式
    for tbl in doc.tables:
        tbl.width = Inches(params['tbl_width'])
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p.style.name != "Normal":
                        skipped.add(f"表格内：{p.style.name}")
                        continue
                    for run in p.runs:
                        set_font(run, params['tbl_cz_font_name'], params['tbl_font_name'])
                        run.font.size = Pt(params['tbl_font_size'])
                    p.paragraph_format.space_before = Pt(params['tbl_space_before'])
                    p.paragraph_format.space_after = Pt(params['tbl_space_after'])
                    p.paragraph_format.line_spacing = params['tbl_line_spacing']
    
    if skipped:
        st.warning(f"跳过样式: {', '.join(sorted(skipped))}")
    
    # 保存到buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# 页脚
st.markdown("---")
st.caption("© 2024 Word自动排版工具 | 专业排版 • 高效便捷")
